from __future__ import annotations

# Завантажувачі різних типів файлів для отримання текстового контенту
import io
import platform
from pathlib import Path
from typing import Tuple, List

# Налаштування шляху Tesseract для Windows
if platform.system() == 'Windows':
    try:
        import pytesseract
        # Спробувати звичайні шляхи встановлення Tesseract
        possible_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        ]
        for path in possible_paths:
            if Path(path).exists():
                pytesseract.pytesseract.tesseract_cmd = path
                print(f"[INFO] Tesseract found at: {path}", flush=True)
                break
        else:
            print("[WARNING] Tesseract OCR not found in standard Windows paths. OCR may not work.", flush=True)
    except ImportError:
        print("[WARNING] pytesseract not available. OCR will not work.", flush=True)


def _load_txt(path: Path) -> str:
    """Просте читання .txt як UTF-8."""
    return Path(path).read_text(encoding="utf-8", errors="ignore")


def _load_pdf(path: Path) -> str:
    """Отримання тексту з PDF за допомогою PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except Exception as exc:
        raise RuntimeError("PyMuPDF (fitz) is required to read PDFs. Add 'pymupdf' to requirements.") from exc

    text_parts = []
    with fitz.open(path) as doc:
        for page in doc:
            # Try multiple text extraction methods for better results
            text = page.get_text()
            if not text.strip():
                # Try alternative extraction method
                text = page.get_text("text")
            if not text.strip():
                # Try with layout preservation
                text = page.get_text("dict")
                if text and "blocks" in text:
                    text_content = []
                    for block in text["blocks"]:
                        if "lines" in block:
                            for line in block["lines"]:
                                if "spans" in line:
                                    for span in line["spans"]:
                                        if "text" in span:
                                            text_content.append(span["text"])
                    text = " ".join(text_content)
            text_parts.append(text)
    return "\n".join(text_parts)


def _extract_pdf_tables(path: Path) -> str:
    """Спроба витягти таблиці з PDF через camelot або tabula-py (якщо доступні)."""
    tables_md: List[str] = []
    try:
        import camelot  # type: ignore
        tables = camelot.read_pdf(str(path), pages="all")
        for t in tables:
            try:
                df = t.df
                tables_md.append(df.to_markdown(index=False))
            except Exception:
                continue
    except Exception:
        try:
            import tabula  # type: ignore
            dfs = tabula.read_pdf(str(path), pages="all", multiple_tables=True)
            for df in dfs or []:
                try:
                    tables_md.append(df.to_markdown(index=False))
                except Exception:
                    continue
        except Exception:
            pass
    return "\n\n".join(tables_md)


def _extract_pdf_images_ocr(path: Path) -> str:
    """OCR тексту зі зображень усередині PDF через PyMuPDF (fitz) + pytesseract."""
    try:
        import fitz
        from PIL import Image
        import pytesseract
    except Exception as exc:
        raise RuntimeError("PyMuPDF, pillow, pytesseract required for PDF OCR.") from exc

    texts: List[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            for img in page.get_images(full=True):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image.get("image")
                if not image_bytes:
                    continue
                try:
                    pil_img = Image.open(io.BytesIO(image_bytes))
                    texts.append(pytesseract.image_to_string(pil_img))
                except Exception:
                    continue
    return "\n".join([t.strip() for t in texts if t and t.strip()])


def _load_docx(path: Path) -> str:
    """Читання абзаців із DOCX через python-docx."""
    try:
        import docx  # python-docx
    except Exception as exc:
        raise RuntimeError("python-docx is required to read DOCX files.") from exc

    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs]
    return "\n".join(paragraphs)


def _extract_docx_tables_markdown(path: Path) -> str:
    """Витягує таблиці з DOCX і повертає markdown-подання."""
    try:
        import docx  # python-docx
    except Exception as exc:
        raise RuntimeError("python-docx is required to read DOCX files.") from exc

    doc = docx.Document(str(path))
    tables_md: List[str] = []
    for table in doc.tables:
        rows: List[List[str]] = []
        for row in table.rows:
            cells = [c.text.replace("\n", " ").strip() for c in row.cells]
            rows.append(cells)
        if not rows:
            continue
        header = rows[0]
        body = rows[1:] if len(rows) > 1 else []
        md = []
        md.append("| " + " | ".join(header) + " |")
        md.append("| " + " | ".join(["---"] * len(header)) + " |")
        for r in body:
            md.append("| " + " | ".join(r) + " |")
        tables_md.append("\n".join(md))
    return "\n\n".join(tables_md)


def _extract_docx_images_ocr(path: Path) -> str:
    """Витягує зображення з DOCX і робить OCR (потрібні pillow+pytesseract)."""
    try:
        import docx
        from PIL import Image
        import pytesseract
    except Exception as exc:
        raise RuntimeError("For DOCX image OCR install python-docx, pillow, pytesseract.") from exc

    doc = docx.Document(str(path))
    texts: List[str] = []
    import io as _io
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_data = rel.target_part.blob
            try:
                img = Image.open(_io.BytesIO(image_data))
                texts.append(pytesseract.image_to_string(img))
            except Exception:
                continue
    return "\n".join([t.strip() for t in texts if t and t.strip()])


def _load_table(path: Path) -> str:
    """Парсинг CSV/XLSX у pandas та перетворення прев'ю у markdown-текст."""
    try:
        import pandas as pd
    except Exception as exc:
        raise RuntimeError("pandas is required to read CSV/XLSX files.") from exc

    if path.suffix.lower() == ".csv":
        df = pd.read_csv(path)
    elif path.suffix.lower() in {".xlsx", ".xlsm"}:
        # Явно вказуємо engine для Excel файлів
        df = pd.read_excel(path, engine='openpyxl')
    elif path.suffix.lower() == ".xls":
        # Для старих Excel файлів
        df = pd.read_excel(path, engine='xlrd')
    else:
        # Fallback для інших форматів
        df = pd.read_excel(path)
    # Convert a sample/preview into text context
    preview = df.head(50)
    return preview.to_markdown(index=False)


def _load_image_ocr(path: Path) -> str:
    """OCR тексту із зображення за допомогою pytesseract + Pillow."""
    try:
        from PIL import Image
        import pytesseract
    except Exception as exc:
        raise RuntimeError("pytesseract and pillow are required for OCR. Install tesseract-ocr separately.") from exc

    image = Image.open(path)
    
    # Configure OCR for better Ukrainian text recognition
    try:
        # Try Ukrainian language first, fallback to English
        raw_text = pytesseract.image_to_string(
            image, 
            lang='ukr+eng',  # Ukrainian + English
            config='--psm 6 --oem 3'  # Assume uniform block of text, use LSTM OCR engine
        )
    except Exception:
        # Fallback to English only if Ukrainian fails
        try:
            raw_text = pytesseract.image_to_string(
                image, 
                lang='eng',
                config='--psm 6 --oem 3'
            )
        except Exception:
            # Final fallback without language specification
            raw_text = pytesseract.image_to_string(image)
    
    # Post-process OCR results to fix common recognition errors
    print(f"[INFO] OCR raw result length: {len(raw_text)} characters", flush=True)
    
    cleaned_text = _fix_ocr_errors(raw_text)
    print(f"[INFO] OCR cleaned result length: {len(cleaned_text)} characters", flush=True)
    
    return cleaned_text


def _fix_ocr_errors(text: str) -> str:
    """Fix common OCR recognition errors for Ukrainian and English text."""
    if not text:
        return text
    
    # Basic cleanup first
    import re
    corrected_text = text
    
    # Fix common spacing issues
    corrected_text = re.sub(r'\s+', ' ', corrected_text)  # Multiple spaces to single
    corrected_text = re.sub(r'([.!?])\s*([A-ZА-Я])', r'\1 \2', corrected_text)  # Space after punctuation
    corrected_text = re.sub(r'([a-zа-я])([A-ZА-Я])', r'\1 \2', corrected_text)  # Space between words
    
    # Apply only the most critical corrections to avoid issues
    critical_corrections = {
        '3Gepirac': 'Результати',
        'PesyIbTaTM': 'Результати',
        'PeslOMYBAHA': 'Резюмування',
        'Kem': 'Key',
        'Cucrema': 'Система',
        'ClicTema': 'Система',
        'cucTema': 'система',
    }
    
    for error, correction in critical_corrections.items():
        corrected_text = corrected_text.replace(error, correction)
    
    return corrected_text.strip()


def load_any(path: Path, force_ocr: bool = False, parse_docx_tables: bool = False, ocr_docx_images: bool = False, parse_pdf_tables: bool = False, ocr_pdf_images: bool = False) -> Tuple[str, str]:
    """Визначає тип файлу за розширенням і повертає текст та позначку для таблиць.

    Автоматично обробляє PDF/DOCX: витягує таблиці та робить OCR зображень за замовчуванням.
    Для зображень (.jpg/.jpeg/.png) OCR виконується автоматично (force_ocr ігнорується).
    
    Підтримувані формати:
    - Текстові: .txt
    - Документи: .pdf, .docx (з таблицями та зображеннями)
    - Таблиці: .csv, .xlsx, .xlsm, .xls
    - Зображення: .jpg, .jpeg, .png (з OCR)
    """
    suffix = path.suffix.lower()
    tables_note = ""

    if suffix in {".txt"}:
        text = _load_txt(path)
    elif suffix in {".pdf"}:
        text_main = _load_pdf(path)
        extras: List[str] = []
        # За замовчуванням вмикаємо витяг таблиць та OCR для PDF, якщо явно не вимкнено
        if parse_pdf_tables or parse_pdf_tables is False or parse_pdf_tables is True:
            # Виконуємо за замовчуванням
            try:
                extra_tables = _extract_pdf_tables(path)
                if extra_tables:
                    extras.append(extra_tables)
            except Exception:
                pass
        try:
            extra_ocr = _extract_pdf_images_ocr(path)
            if extra_ocr:
                extras.append(extra_ocr)
        except Exception:
            pass
        text = "\n\n".join([t for t in [text_main] + extras if t])
    elif suffix in {".docx"}:
        text_main = _load_docx(path)
        extras: List[str] = []
        # За замовчуванням додаємо таблиці та OCR зображень
        try:
            md = _extract_docx_tables_markdown(path)
            if md:
                extras.append(md)
                print(f"[INFO] Extracted {len(md.split(chr(10)))} lines of table data from DOCX", flush=True)
        except Exception as e:
            print(f"[WARNING] Failed to extract tables from DOCX: {e}", flush=True)
        try:
            ocr_text = _extract_docx_images_ocr(path)
            if ocr_text:
                extras.append(ocr_text)
                print(f"[INFO] Extracted {len(ocr_text)} characters from DOCX images", flush=True)
        except Exception as e:
            print(f"[WARNING] Failed to OCR images from DOCX: {e}", flush=True)
        text = "\n\n".join([t for t in [text_main] + extras if t])
    elif suffix in {".csv", ".xlsx", ".xlsm", ".xls"}:
        text = _load_table(path)
        tables_note = "Tabular data preview included in context."
        print(f"[INFO] Loaded table from {suffix} file: {len(text)} characters", flush=True)
    elif suffix in {".jpg", ".jpeg", ".png"}:
        # OCR завжди для зображень
        text = _load_image_ocr(path)
    else:
        # Якщо розширення відсутнє або невідоме — пробуємо читати як звичайний текст
        try:
            text = _load_txt(path)
        except Exception as exc:
            raise ValueError(f"Unsupported file type: {suffix}") from exc

    return text, tables_note


